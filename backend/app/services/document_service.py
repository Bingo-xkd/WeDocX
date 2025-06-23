"""
文档服务模块，提供文档格式转换和保存功能
"""

import os
import re
from typing import List, Optional, Tuple

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor


def _clean_text(text: str) -> str:
    """清理文本内容，保留换行符，处理特殊字符和多余空白"""
    # 先替换多余的空格，但保留换行符
    text = re.sub(r"[ \t\u3000]+", " ", text)  # 只替换空格和制表符，不动\n
    # 清理特殊字符但保留中文、英文、标点和换行
    text = re.sub(r'[^\w\s\u4e00-\u9fff,.?!，。？！、:：;；""' "()（）\n-]", "", text)
    # 去除每行首尾空白
    text = "\n".join(line.strip() for line in text.splitlines())
    return text.strip()


def _extract_text_with_structure(elem: Tag) -> List[Tuple[str, str, str]]:
    """
    提取HTML元素中的文本，并保持结构
    返回格式：List of (text, tag_name, class_name)
    """
    results = []

    if isinstance(elem, NavigableString):
        text = _clean_text(str(elem))
        if text:
            results.append((text, "", ""))
        return results

    if elem.name in ["script", "style", "meta", "link"]:
        return results

    # 处理特殊标签
    if elem.name == "br":
        results.append(("\n", "br", ""))
        return results

    # 获取元素的class
    class_name = " ".join(elem.get("class", []))

    # 处理列表
    if elem.name in ["ul", "ol"]:
        for i, li in enumerate(elem.find_all("li", recursive=False)):
            prefix = f"{i+1}. " if elem.name == "ol" else "• "
            text = _clean_text(li.get_text())
            if text:
                results.append((f"{prefix}{text}", "li", class_name))
        results.append(("\n", "ulol", class_name))
        return results

    # 处理表格
    if elem.name == "table":
        for row in elem.find_all("tr", recursive=False):
            cells = []
            for cell in row.find_all(["td", "th"], recursive=False):
                cells.append(_clean_text(cell.get_text()))
            if cells:
                results.append((" | ".join(cells), "table", class_name))
        results.append(("\n", "table", ""))
        return results

    # 递归处理子元素
    for child in elem.children:
        if isinstance(child, (Tag, NavigableString)):
            results.extend(_extract_text_with_structure(child))

    # 为块级元素添加换行
    if elem.name in ["p", "div", "h1", "h2", "h3", "h4", "h5", "h6"]:
        if results and not results[-1][0].endswith("\n"):
            results.append(("\n", "", ""))
        if results and not results[0][0].startswith("\n"):
            results.insert(0, ("\n", "", ""))

    return results


def convert_html_to_txt(
    html_content: str, output_path: str, title: Optional[str] = None
) -> str:
    """
    将HTML内容转换并保存为TXT文件，保持基本结构和格式。

    Args:
        html_content: HTML内容
        output_path: 输出文件路径
        title: 可选的文档标题

    Returns:
        str: 保存的文件路径
    """
    soup = BeautifulSoup(html_content, "html.parser")

    with open(output_path, "w", encoding="utf-8") as f:
        # 写入标题
        if title:
            f.write(f"{title}\n{'='*len(title)}\n\n")

        # 提取并写入结构化文本
        content = []
        for text, tag, _ in _extract_text_with_structure(soup.body or soup):
            content.append(text)

        # 处理连续换行
        text_content = "".join(content)
        text_content = re.sub(r"\n{3,}", "\n\n", text_content)
        f.write(text_content)

    return output_path


def convert_html_to_docx(
    html_content: str, output_path: str, title: Optional[str] = None
) -> str:
    """
    将HTML内容转换并保存为Word(docx)文件，保持文档结构和基本样式。

    Args:
        html_content: HTML内容
        output_path: 输出文件路径
        title: 可选的文档标题

    Returns:
        str: 保存的文件路径
    """
    soup = BeautifulSoup(html_content, "html.parser")
    doc = Document()

    # 设置标题
    if title:
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()  # 添加空行

    def get_heading_level(tag_name: str) -> int:
        """获取标题级别"""
        if not tag_name.startswith("h"):
            return 0
        level = int(tag_name[1])
        return min(level, 9)  # Word支持最多9级标题

    # 遍历并处理结构化内容
    current_paragraph = None

    for text, tag, classes in _extract_text_with_structure(soup.body or soup):
        # 处理标题
        if tag.startswith("h") and tag != "hr":
            level = get_heading_level(tag)
            if text.strip():
                doc.add_heading(text.strip(), level=level)
                current_paragraph = None
            continue

        # 处理列表和表格
        if tag in ["li", "table"]:
            if text.strip():
                current_paragraph = doc.add_paragraph(text.strip())
                current_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            continue

        # 处理普通段落
        if tag in ["p", "div"] or not tag:
            if text.strip():
                if not current_paragraph:
                    current_paragraph = doc.add_paragraph()
                current_paragraph.add_run(text.strip())
                if text.endswith("\n"):
                    current_paragraph = None
            continue

        # 处理换行
        if tag == "br":
            current_paragraph = None
            continue

    # 保存文档
    doc.save(output_path)
    return output_path
