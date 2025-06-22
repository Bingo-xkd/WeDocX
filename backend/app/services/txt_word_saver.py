import os
from typing import Optional
from docx import Document
from bs4 import BeautifulSoup

def save_as_txt(txt_path: str, text: str, title: Optional[str] = None):
    """
    保存纯文本内容为txt文件。
    :param txt_path: txt文件保存路径
    :param text: 纯文本内容
    :param title: 可选，文档标题
    """
    with open(txt_path, 'w', encoding='utf-8') as f:
        if title:
            f.write(title + '\n\n')
        f.write(text)

def save_as_docx(docx_path: str, html: str, title: Optional[str] = None):
    """
    保存HTML内容为Word(docx)文件，仅保留文本和基本结构。
    :param docx_path: docx文件保存路径
    :param html: 网页HTML内容
    :param title: 可选，文档标题
    """
    soup = BeautifulSoup(html, 'html.parser')
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    # 提取正文内容
    body = soup.body
    if body:
        for elem in body.descendants:
            if elem.name == 'p' and elem.get_text(strip=True):
                doc.add_paragraph(elem.get_text(strip=True))
            elif elem.name in ['h1', 'h2', 'h3']:
                doc.add_heading(elem.get_text(strip=True), level=2)
    else:
        doc.add_paragraph(soup.get_text())
    doc.save(docx_path) 